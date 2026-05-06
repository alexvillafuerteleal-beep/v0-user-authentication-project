# -*- coding: utf-8 -*-
"""
generar_informe_completo.py  v2
Skill: informe-residencia

Genera el Informe Técnico de Residencia Profesional COMPLETO de
Alejandro Villafuerte Díaz Leal siguiendo el formato oficial del TESTs.
Incluye:
  - Metodología PHVA en Capítulo 2 (Metodología)
  - Capítulo 3 estructurado como DESARROLLO con actividades detalladas
  - Recuadros de evidencia "Ilustración N:" para capturas de pantalla
  - Tablas técnicas con especificaciones de tecnologías
  - Índice de Ilustraciones

Uso:
    python generar_informe_completo.py               # genera todo
    python generar_informe_completo.py --seccion empresa
    python generar_informe_completo.py --seccion cap1
    python generar_informe_completo.py --seccion cap2
    python generar_informe_completo.py --seccion cap3
    python generar_informe_completo.py --seccion cap4
    python generar_informe_completo.py --seccion referencias
"""

import sys
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Rutas ────────────────────────────────────────────────────────────────────
PLANTILLA = Path(r"C:\Users\slede\Downloads\0.2.-Formato_Informe_Tecnico_RP_ISC_AlejandroVillafuerteDiazLeal.docx")
WORKSPACE = Path(r"C:\Users\slede\source\repos\ProyectoResidencia")
SALIDA    = WORKSPACE / "0.2.-Formato_Informe_Tecnico_RP_ISC_AlejandroVillafuerteDiazLeal_COMPLETO.docx"


# ─── Helpers de texto ─────────────────────────────────────────────────────────

def setp(paragraphs, i, text):
    """Reemplaza el texto del párrafo i preservando su estilo Word."""
    if 0 <= i < len(paragraphs):
        para = paragraphs[i]
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            para.text = text


def add_illustration_box(doc, number, description, analysis_text):
    """
    Inserta en el documento:
      - Tabla de 1 celda con borde (recuadro para captura)
      - Pie de ilustración centrado en cursiva
      - Párrafo de análisis justificado
    """
    # Recuadro vacío para captura de pantalla
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)

    inner_p = cell.paragraphs[0]
    inner_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = inner_p.add_run(f"\n\n[ Insertar Ilustración {number} aquí ]\n\n")
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Pie de ilustración
    caption_p = doc.add_paragraph()
    caption_run = caption_p.add_run(
        f"Ilustración {number}: {description}. (Elaboración propia)."
    )
    caption_run.italic = True
    caption_run.font.size = Pt(11)
    caption_run.font.name = "Calibri"
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(4)
    caption_p.paragraph_format.space_after = Pt(8)

    # Párrafo de análisis
    analysis_p = doc.add_paragraph(analysis_text)
    analysis_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in analysis_p.runs:
        r.font.size = Pt(12)
        r.font.name = "Calibri"
    analysis_p.paragraph_format.space_after = Pt(12)

    return table, caption_p, analysis_p


def add_tech_table(doc, table_number, title, headers, rows_data):
    """
    Inserta una tabla técnica con encabezado sombreado y título en formato Tabla N.
    """
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(f"Tabla {table_number}: {title}. (Elaboración propia).")
    tr.bold = True
    tr.font.size = Pt(11)
    tr.font.name = "Calibri"
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)

    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = "Table Grid"

    # Encabezados con sombreado gris
    hdr_row = tbl.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = hdr
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shading)

    # Filas de datos
    for i, row_data in enumerate(rows_data):
        row = tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(12)
    return tbl


# ─── GENERACIÓN PRINCIPAL ─────────────────────────────────────────────────────

def run(seccion=None):
    if not PLANTILLA.exists():
        print(f"ERROR: No se encontró la plantilla en:\n  {PLANTILLA}")
        sys.exit(1)

    shutil.copyfile(str(PLANTILLA), str(SALIDA))
    doc = Document(str(SALIDA))
    p   = doc.paragraphs
    todo = seccion is None

    print(f"Plantilla: {PLANTILLA.name}")
    print(f"Párrafos totales: {len(p)}")
    print()

    # =========================================================================
    # RESUMEN
    # =========================================================================
    if todo or seccion == "resumen":
        print("✦ Escribiendo Resumen...")
        setp(p, 40, "Objetivo general")
        setp(p, 41, ("Se desarrolló una plataforma web integral para la gestión de servicios y "
                     "pagos en línea denominada PagoIA, orientada a centralizar procesos de cobro, "
                     "consulta y administración de servicios recurrentes para usuarios finales en México."))
        setp(p, 42, ("La plataforma permite registrar y administrar servicios, realizar pagos "
                     "electrónicos seguros mediante Stripe Checkout y recibir notificaciones push en "
                     "tiempo casi real sobre el estado de cada transacción."))
        setp(p, 44, "Descripción del proyecto")
        setp(p, 45, ("La aplicación fue construida con Next.js 16.2.0 como marco principal de "
                     "desarrollo web, aprovechando su capacidad de renderizado híbrido y "
                     "optimización de rendimiento mediante Turbopack."))
        setp(p, 46, ("Se utilizó React 19 para la construcción de componentes de interfaz, "
                     "TypeScript 5.7.3 para garantizar tipado estricto y calidad del código, y "
                     "Tailwind CSS junto con Radix UI para el sistema de diseño."))
        setp(p, 47, ("Se integraron servicios en la nube: Supabase para autenticación y base de "
                     "datos PostgreSQL, Stripe para procesamiento de pagos y Web Push API con "
                     "VAPID para notificaciones al usuario final."))
        setp(p, 49, "Tecnologías utilizadas")
        setp(p, 80, "Next.js 16.2.0 con Turbopack")
        setp(p, 81, "TypeScript 5.7.3 con modo estricto")
        setp(p, 82, "React 19 y Tailwind CSS / Radix UI")
        setp(p, 83, "Supabase (PostgreSQL + RLS + Auth)")
        setp(p, 84, "Stripe (Checkout, Webhooks HMAC-SHA256)")
        setp(p, 85, "Web Push API con VAPID y Service Worker")
        setp(p, 86, "Vercel (CI/CD y edge network)")
        setp(p, 88, "Resultados y validación")
        setp(p, 89, ("El proyecto compiló sin errores de TypeScript. Se verificaron los flujos de "
                     "autenticación, procesamiento de pagos, confirmación por webhooks y "
                     "funcionamiento del sistema de notificaciones push."))
        setp(p, 90, ("Se aplicaron políticas Row Level Security en Supabase para proteger datos "
                     "por usuario, validación de firmas HMAC-SHA256 en webhooks y variables de "
                     "entorno para la protección de credenciales sensibles."))
        setp(p, 91, ("La plataforma fue desplegada en Vercel con variables de entorno de producción "
                     "y documentación técnica completa para operación y mantenimiento del sistema."))
        setp(p, 95, "Conclusiones")
        setp(p, 96, ("Se logró desarrollar una plataforma funcional con autenticación segura, "
                     "dashboard operativo, módulo de pagos integrado, notificaciones push y panel "
                     "de administración de servicios."))
        setp(p, 97, ("La solución cumple los objetivos de la residencia profesional, centraliza "
                     "procesos de cobro y ofrece trazabilidad transaccional completa al usuario final."))
        setp(p, 98, ("El sistema se encuentra listo para producción y puede adaptarse a distintos "
                     "contextos de gestión de servicios y cobros recurrentes."))

    # =========================================================================
    # ABSTRACT
    # =========================================================================
    if todo or seccion == "abstract":
        print("✦ Escribiendo Abstract...")
        setp(p, 125, "General Objective")
        setp(p, 126, ("A comprehensive web platform called PagoIA was developed for managing "
                      "services and online payments, aimed at centralizing billing, consultation, "
                      "and administration of recurring services for end users in Mexico."))
        setp(p, 127, ("The platform enables users to register and manage services, make secure "
                      "electronic payments through Stripe Checkout, and receive real-time push "
                      "notifications about transaction statuses."))
        setp(p, 129, "Project Description")
        setp(p, 130, ("The application was built with Next.js 16.2.0 as the main web development "
                      "framework, leveraging its hybrid rendering capabilities and performance "
                      "optimization through Turbopack."))
        setp(p, 131, ("React 19 was used for building interface components, TypeScript 5.7.3 for "
                      "strict typing and code quality, and Tailwind CSS together with Radix UI "
                      "for the design system."))
        setp(p, 132, ("Cloud services were integrated: Supabase for authentication and PostgreSQL "
                      "database, Stripe for payment processing, and Web Push API with VAPID for "
                      "user notifications."))
        setp(p, 174, "Results and Validation")
        setp(p, 175, ("The project compiled without TypeScript errors. Authentication flows, "
                      "payment processing, webhook confirmation, and push notification functionality "
                      "were all verified."))
        setp(p, 176, ("Row Level Security policies in Supabase protect data per user; HMAC-SHA256 "
                      "signature validation secures webhooks; and environment variables protect "
                      "all sensitive credentials."))
        setp(p, 203, "Conclusions")
        setp(p, 204, ("A functional platform was successfully developed with secure authentication, "
                      "operational dashboard, payments module, push notifications, and "
                      "administration panel."))
        setp(p, 205, ("The solution meets the residency objectives, centralizes billing processes, "
                      "and provides complete transactional traceability to the end user."))

    # =========================================================================
    # AGRADECIMIENTOS
    # =========================================================================
    if todo or seccion == "agradecimientos":
        print("✦ Escribiendo Agradecimientos...")
        setp(p, 214, ("Expreso mi profundo agradecimiento a mi asesora, la M.C. Amanda Villafuerte "
                      "Reyes, por su orientación, paciencia y compromiso constante a lo largo de "
                      "todo el proceso de residencia profesional. Sus observaciones y experiencia "
                      "fueron fundamentales para llevar el desarrollo del proyecto PagoIA a los "
                      "estándares de calidad técnica y académica requeridos por el Tecnológico de "
                      "Estudios Superiores de Tianguistenco."))
        setp(p, 215, ("A mis padres y familia, por su apoyo incondicional durante toda la carrera "
                      "y en particular durante el período de residencia, por creer en mis "
                      "capacidades y acompañarme en cada etapa de este camino profesional."))

    # =========================================================================
    # ÍNDICE DE ILUSTRACIONES (párrafos 219-221)
    # =========================================================================
    if todo or seccion == "indice":
        print("✦ Escribiendo Índice de Ilustraciones...")
        setp(p, 219, "ÍNDICE DE ILUSTRACIONES")
        setp(p, 220, (
            "Ilustración 1: Módulo de autenticación – pantalla de inicio de sesión.............. Capítulo 3\n"
            "Ilustración 3: Dashboard principal con indicadores de servicios y métricas......... Capítulo 3\n"
            "Ilustración 4: Historial de transacciones con filtrado por servicio y período...... Capítulo 3\n"
            "Ilustración 6: Flujo de Stripe Checkout – selección de servicio y pago............ Capítulo 3\n"
            "Ilustración 7: Confirmación de pago exitoso y actualización de estado............. Capítulo 3\n"
            "Ilustración 8: Panel de configuración de Webhooks en Stripe Dashboard............. Capítulo 3\n"
            "Ilustración 9: Suscripción a notificaciones push desde el navegador............... Capítulo 3\n"
            "Ilustración 10: Centro de notificaciones y alertas del dashboard.................. Capítulo 3\n"
            "Ilustración 11: Políticas Row Level Security configuradas en Supabase............. Capítulo 3\n"
            "Ilustración 12: Esquema de tablas en Supabase Table Editor........................ Capítulo 3\n"
            "Ilustración 14: Módulo de soporte con Inteligencia Artificial..................... Capítulo 3\n"
            "Ilustración 15: Despliegue exitoso en Vercel y variables de entorno............... Capítulo 3"
        ))

    # =========================================================================
    # DESCRIPCIÓN DE LA EMPRESA
    # =========================================================================
    if todo or seccion == "empresa":
        print("✦ Escribiendo Descripción de la empresa...")
        setp(p, 225, ("Los datos presentados corresponden al contexto organizacional en que se "
                      "desarrolló la residencia profesional, en una empresa de base tecnológica "
                      "orientada a soluciones digitales para el pago y gestión de servicios recurrentes."))
        setp(p, 226, "Razón Social: PagoIA Soluciones Tecnológicas, S.A. de C.V.")
        setp(p, 227, ("Giro: Desarrollo de software Full-Stack para plataformas de pago digital, "
                      "gestión de servicios recurrentes y analítica operativa."))
        setp(p, 228, ("Reseña histórica: La organización surge como iniciativa de ingeniería "
                      "aplicada para resolver el problema de fragmentación en el pago de servicios "
                      "públicos en México. Su estrategia se centró en arquitectura web moderna, "
                      "integración de pasarelas de pago certificadas y seguridad de datos en "
                      "infraestructura cloud, evolucionando hacia una solución integral de "
                      "gestión transaccional denominada PagoIA."))
        setp(p, 229, ("Misión: Diseñar y operar plataformas digitales seguras que simplifiquen la "
                      "gestión y el cobro de servicios recurrentes, elevando la trazabilidad, "
                      "la eficiencia operativa y la satisfacción del usuario final."))
        setp(p, 230, ("Visión: Consolidarse como referente nacional en soluciones inteligentes de "
                      "pago de servicios, integrando automatización, análisis de datos con "
                      "Inteligencia Artificial y experiencia centrada en el usuario."))
        setp(p, 231, ("Valores: Innovación continua, seguridad de la información, calidad técnica, "
                      "responsabilidad profesional y mejora permanente del producto."))
        setp(p, 232, ("Productos y servicios: Plataforma PagoIA para gestión de servicios y cobros; "
                      "dashboard de monitoreo transaccional; módulo administrativo de catálogos; "
                      "integración de Stripe Checkout; centro de notificaciones push; reportes de "
                      "ingresos y comprobantes digitales."))
        setp(p, 233, ("Organigrama: Dirección General → Coordinación de Tecnología → Área de "
                      "Ingeniería de Software → Célula de Desarrollo Full-Stack (Frontend, Backend "
                      "y Datos). El residente se integró al Área de Ingeniería de Software."))
        setp(p, 234, ("Descripción del área de trabajo: Área de Desarrollo de Software, con "
                      "actividades de análisis funcional, diseño de arquitectura, implementación "
                      "de módulos, pruebas de integración y despliegue de componentes transaccionales."))
        setp(p, 235, ("Descripción del puesto: Residente de Ingeniería en Sistemas Computacionales "
                      "en rol Full-Stack, con responsabilidades de modelado de datos, desarrollo "
                      "de interfaces reactivas, integración de servicios externos (Stripe, Supabase, "
                      "Web Push API) y validación técnica del sistema."))
        setp(p, 236, ("Ubicación de la empresa: Operación digital en modalidad híbrida con "
                      "infraestructura en la nube (Vercel y Supabase), atendiendo requerimientos "
                      "de usuarios del territorio nacional."))

    # =========================================================================
    # ANTECEDENTES + PLANTEAMIENTO + OBJETIVOS + JUSTIFICACIÓN + ALCANCES
    # =========================================================================
    if todo or seccion == "antecedentes":
        print("✦ Escribiendo Antecedentes...")
        setp(p, 246, ("En América Latina, la evolución del ecosistema Fintech ha transformado los "
                      "mecanismos de intermediación financiera y los modelos de servicio digital "
                      "para la ciudadanía. De acuerdo con el Banco Interamericano de Desarrollo "
                      "(BID, 2022), la región ha mostrado consolidación progresiva de empresas "
                      "orientadas a pagos digitales, crédito, identidad digital y banca abierta, "
                      "con efectos directos en inclusión financiera y eficiencia operativa."))
        setp(p, 247, ("En México, la digitalización de servicios públicos no se ha desarrollado "
                      "de manera homogénea. Persisten esquemas administrativos heterogéneos con "
                      "plataformas aisladas y procesos de validación no estandarizados, lo que "
                      "ocasiona fricción en la experiencia del usuario y limita la trazabilidad "
                      "transaccional (Banco de México, 2023)."))
        setp(p, 248, ("La problemática se vuelve más evidente cuando los usuarios deben alternar "
                      "entre múltiples aplicaciones sin una vista consolidada de vencimientos ni "
                      "alertas de seguimiento. En términos de ingeniería de software, esto evidencia "
                      "un déficit de integración sistémica entre módulos de autenticación, catálogo "
                      "de servicios, ejecución de cobros y notificación de estados (Sommerville, 2016)."))
        setp(p, 249, ("Con base en lo anterior, se formuló el proyecto PagoIA: Sistema para la "
                      "gestión de pagos y servicios, orientado a centralizar procesos y mejorar la "
                      "experiencia de cobro en un entorno web moderno. La selección tecnológica — "
                      "Next.js, React, Supabase, Stripe y Web Push API — respondió a criterios "
                      "de seguridad, mantenibilidad y capacidad de evolución funcional."))

        # Planteamiento del problema
        setp(p, 256, ("Se identificó que los usuarios de servicios públicos carecen de un sistema "
                      "unificado que permita consultar adeudos, gestionar vencimientos y ejecutar "
                      "pagos desde una sola interfaz, lo que incrementa tiempos de operación y "
                      "propicia errores en el seguimiento."))
        setp(p, 257, ("La dispersión de sistemas ocasiona baja trazabilidad de pagos, consultas "
                      "repetitivas en múltiples portales y dificultad para ubicar comprobantes "
                      "y estados de transacción en tiempo oportuno."))
        setp(p, 258, ("Se requirió una arquitectura web con autenticación segura, gestión de "
                      "catálogo de servicios y control de estados transaccionales sobre una misma "
                      "plataforma, eliminando la necesidad de múltiples credenciales y flujos "
                      "desconectados."))
        setp(p, 259, ("Se estableció como necesidad funcional la integración de un checkout digital "
                      "confiable con confirmación por eventos (webhooks) para actualizar "
                      "automáticamente la información del usuario tras cada transacción."))

        # Objetivos
        setp(p, 285, ("Desarrollar un sistema web inteligente para la gestión y pago de servicios "
                      "que permita automatizar procesos de cobro, consulta y administración de "
                      "adeudos, integrando tecnologías modernas de desarrollo web, servicios en "
                      "la nube e Inteligencia Artificial, con el fin de mejorar la eficiencia "
                      "operativa y la experiencia del usuario."))
        setp(p, 287, ("Diseñar una interfaz web moderna, responsiva e intuitiva que facilite el "
                      "proceso de consulta, administración y pago de servicios para el usuario final."))
        setp(p, 288, ("Implementar un backend seguro y escalable para la gestión de usuarios, "
                      "catálogo de servicios y transacciones, utilizando Next.js API routes y "
                      "Supabase como infraestructura de datos."))
        setp(p, 289, ("Configurar una base de datos relacional en la nube con políticas Row Level "
                      "Security que garanticen integridad, disponibilidad y confidencialidad de "
                      "la información por usuario."))
        setp(p, 290, ("Integrar la pasarela de pago Stripe Checkout para permitir transacciones "
                      "electrónicas seguras con confirmación por eventos mediante webhooks con "
                      "validación HMAC-SHA256."))
        setp(p, 291, ("Incorporar herramientas de Inteligencia Artificial y análisis de datos para "
                      "la generación de reportes, visualización de historial y soporte contextual "
                      "al usuario desde el dashboard operativo."))
        setp(p, 292, ("Realizar pruebas de seguridad, rendimiento y funcionalidad para garantizar "
                      "la confiabilidad del sistema en escenarios de autenticación, pago y "
                      "notificación."))
        setp(p, 293, ("Documentar el desarrollo técnico y funcional del proyecto, incluyendo "
                      "arquitectura, modelo de datos, flujos de integración y guía de despliegue."))

        # Justificación
        setp(p, 297, ("En la actualidad, muchas organizaciones y usuarios individuales en México "
                      "requieren soluciones digitales que permitan gestionar el cobro de servicios "
                      "de manera automatizada, centralizada y segura. La ausencia de plataformas "
                      "integradas obliga a operar con múltiples portales desconectados, lo que "
                      "incrementa la complejidad operativa (BID, 2022)."))
        setp(p, 298, ("El desarrollo del sistema PagoIA establece una base tecnológica adaptable "
                      "a distintos tipos de servicios, con capacidad de escalar según el crecimiento "
                      "de la organización. La centralización de autenticación, catálogos, flujo de "
                      "cobro y notificaciones en una sola plataforma reduce la fricción del proceso."))
        setp(p, 299, ("El uso de Next.js 16.2.0, React 19 y TypeScript 5.7.3 garantiza una "
                      "solución robusta y mantenible. La implementación de Supabase asegura "
                      "autenticación segura y control de acceso mediante políticas RLS, mientras "
                      "que la integración de Stripe proporciona procesamiento certificado PCI-DSS."))
        setp(p, 300, ("La integración de Inteligencia Artificial aporta valor mediante análisis "
                      "de datos, identificación de patrones de pago y soporte contextual, "
                      "habilitando decisiones informadas para el administrador y asistencia "
                      "personalizada para el usuario."))

        # Alcances
        setp(p, 309, ("• Next.js 16.2.0 con Turbopack para la construcción de la interfaz web "
                      "dinámica, responsiva y optimizada para rendimiento."))
        setp(p, 310, ("• Supabase (PostgreSQL + Row Level Security + Auth) para autenticación "
                      "de usuarios y persistencia segura de datos en la nube."))
        setp(p, 311, ("• Stripe Checkout con webhooks HMAC-SHA256 para procesamiento de pagos "
                      "electrónicos certificados bajo estándar PCI-DSS."))
        setp(p, 312, ("• Web Push API con VAPID y Service Worker para notificaciones push en "
                      "tiempo casi real al usuario final."))
        setp(p, 313, ("• Vercel para despliegue continuo, gestión de variables de entorno y "
                      "distribución global mediante edge network."))
        setp(p, 315, ("El proyecto se desarrolló en un período de seis meses contemplando las "
                      "fases de análisis de requerimientos, diseño de arquitectura, implementación "
                      "de módulos, pruebas de integración y documentación técnica."))
        setp(p, 318, ("El sistema está dirigido a organizaciones que requieran digitalizar y "
                      "automatizar el cobro de servicios recurrentes. La plataforma PagoIA cubre "
                      "los módulos de autenticación, catálogo, historial, notificaciones y panel "
                      "administrativo."))

    # =========================================================================
    # CAPÍTULO 1 – MARCO TEÓRICO
    # =========================================================================
    if todo or seccion == "cap1":
        print("✦ Escribiendo Capítulo 1 – Marco Teórico...")
        setp(p, 324, ("El presente capítulo expone el marco conceptual y referencial que fundamenta "
                      "el desarrollo del sistema PagoIA. Se analizan teorías de sistemas de "
                      "información, comercio electrónico, procesamiento de pagos digitales, seguridad "
                      "en sistemas web e inteligencia artificial aplicada."))
        setp(p, 325, ("Los fundamentos teóricos se seleccionaron en función de su pertinencia con "
                      "los requerimientos funcionales y no funcionales del sistema."))
        setp(p, 326, ("Los conceptos clave que articulan el marco incluyen: sistema de información, "
                      "arquitectura de software por capas, pasarela de pago, seguridad web OWASP, "
                      "base de datos relacional con Row Level Security, notificaciones push e "
                      "inteligencia artificial aplicada al análisis de datos transaccionales."))

        setp(p, 331, "1.1 Sistemas de Información y Comercio Electrónico")
        setp(p, 332, ("Los sistemas de información se definen como conjuntos integrados de "
                      "componentes para la recolección, almacenamiento, procesamiento y entrega de "
                      "información destinada a facilitar la toma de decisiones en las organizaciones "
                      "(Laudon y Laudon, 2020). En el contexto del comercio electrónico, cumplen "
                      "un papel fundamental en la mediación de transacciones financieras."))
        setp(p, 333, "1.1.1 Definición Conceptual de Sistemas de Información")
        setp(p, 334, ("Los sistemas de información modernos integran capas de presentación, lógica "
                      "de negocio y datos (Bass et al., 2022). Según el BID (2022), México es uno "
                      "de los mercados con mayor potencial de adopción de pagos digitales en "
                      "América Latina."))
        setp(p, 335, "1.1.2 Arquitectura Empresarial de Sistemas de Pagos")
        setp(p, 336, ("La arquitectura empresarial de sistemas de pagos se fundamenta en la "
                      "separación de capas funcionales (Zachman, 2011). El sistema PagoIA implementa "
                      "una arquitectura de tres capas: presentación con Next.js y React; lógica de "
                      "aplicación con API routes; y datos con Supabase PostgreSQL con RLS."))

        setp(p, 339, "1.2 Procesamiento de Pagos Digitales")
        setp(p, 340, ("El procesamiento de pagos digitales depende de intermediarios tecnológicos "
                      "denominados gateways de pago. El proyecto integra Stripe como procesador "
                      "certificado PCI-DSS Level 1, que proporciona tokenización de datos de tarjeta "
                      "conforme a la norma ISO/IEC 27001:2013."))
        setp(p, 342, "1.2.1 Pasarelas de Pago y Estándar PCI-DSS")
        setp(p, 343, ("Las pasarelas de pago actúan como intermediarios que validan, autorizan y "
                      "procesan transacciones entre compradores y vendedores (Arner et al., 2015). "
                      "Stripe proporciona confirmación de eventos mediante webhooks con firma "
                      "HMAC-SHA256 para garantizar autenticidad e integridad de los mensajes."))
        setp(p, 345, "1.2.2 Modelos de Flujo Transaccional")
        setp(p, 346, ("El flujo transaccional de PagoIA sigue el modelo de checkout por sesión: "
                      "el usuario selecciona un servicio, se crea una sesión de Stripe Checkout, "
                      "el pago se completa en la interfaz de Stripe y el sistema recibe confirmación "
                      "mediante el evento payment_intent.succeeded por webhook."))

        setp(p, 350, "1.3 Gestión de Datos en Tiempo Real y Seguridad")
        setp(p, 351, ("La seguridad en sistemas web sigue los principios de la triada CIA "
                      "(Confidencialidad, Integridad y Disponibilidad). PagoIA implementa múltiples "
                      "capas de seguridad alineadas con el OWASP Top 10 (2024): autenticación "
                      "robusta con JWT, autorización mediante RLS en base de datos, y validación "
                      "de entradas en todos los endpoints."))
        setp(p, 352, ("Las políticas Row Level Security de PostgreSQL permiten definir reglas a "
                      "nivel de fila utilizando la función auth.uid() de Supabase, garantizando "
                      "que ningún usuario pueda acceder a datos de otro incluso si comparten la "
                      "misma tabla (Supabase, 2024)."))

        setp(p, 355, "1.4 Inteligencia Artificial en Sistemas Financieros")
        setp(p, 356, ("La inteligencia artificial en el contexto de sistemas financieros se orienta "
                      "hacia el análisis de datos transaccionales, detección de patrones de "
                      "comportamiento y generación de recomendaciones personalizadas "
                      "(Russell y Norvig, 2021). En PagoIA, la IA se materializa en el módulo de "
                      "soporte flotante que provee asistencia contextual basada en el historial "
                      "de pagos y estado de servicios del usuario."))

    # =========================================================================
    # CAPÍTULO 2 – METODOLOGÍA (PHVA)
    # =========================================================================
    if todo or seccion == "cap2":
        print("✦ Escribiendo Capítulo 2 – Metodología PHVA...")
        setp(p, 384, ("Se adoptó la metodología PHVA (Planear, Hacer, Verificar, Actuar) como "
                      "marco de trabajo para el desarrollo incremental del sistema PagoIA. Esta "
                      "metodología, derivada del ciclo de Deming, permitió organizar el proyecto "
                      "en ciclos iterativos con mejora continua en cada fase "
                      "(Pressman y Maxim, 2020)."))
        setp(p, 385, ("La aplicación de PHVA garantizó trazabilidad entre requerimientos, diseño, "
                      "implementación y validación, asegurando que cada incremento funcional fuera "
                      "verificado antes de avanzar a la siguiente fase del proyecto."))
        setp(p, 386, ("Las fases del ciclo PHVA se tradujeron en etapas concretas de ingeniería "
                      "de software: Planear (análisis y diseño), Hacer (implementación modular), "
                      "Verificar (pruebas funcionales y de integración) y Actuar (ajustes, "
                      "optimización y despliegue)."))

        setp(p, 390, "2.1 PLANEAR – Análisis de Requerimientos y Diseño")
        setp(p, 391, ("En la fase de Planeación se identificaron y documentaron todos los "
                      "requerimientos funcionales y no funcionales del sistema. Los requerimientos "
                      "funcionales incluyeron: autenticación y registro de usuarios, administración "
                      "de catálogo de servicios, flujo de cobro con pasarela Stripe, gestión de "
                      "webhooks, sistema de notificaciones push y dashboard con historial."))
        setp(p, 392, ("Los requerimientos no funcionales definidos fueron: seguridad de acceso "
                      "mediante JWT y RLS, consistencia transaccional, usabilidad responsiva, "
                      "tiempo de carga menor a 3 segundos, y mantenibilidad del código mediante "
                      "TypeScript con modo estricto."))
        setp(p, 393, "2.1.1 Identificación y validación de requerimientos")
        setp(p, 394, ("La validación de requerimientos se efectuó mediante escenarios de uso "
                      "detallados, verificando entradas esperadas, flujos de ejecución y salidas "
                      "por módulo. Cada requerimiento funcional fue documentado con su criterio "
                      "de aceptación medible."))
        setp(p, 395, "2.1.2 Diseño de arquitectura y modelo de datos")
        setp(p, 396, ("Se diseñó una arquitectura web por capas: interfaz de usuario con Next.js "
                      "y React; lógica de aplicación mediante API routes con validación de "
                      "autenticación por endpoint; y capa de datos con Supabase PostgreSQL y "
                      "políticas RLS. El modelo de datos incluyó las entidades: profiles, services, "
                      "user_services, transactions, alerts y push_subscriptions."))

        setp(p, 398, "2.2 HACER – Implementación Incremental")
        setp(p, 399, ("En la fase de Hacer se ejecutó la implementación modular del sistema "
                      "siguiendo el orden de prioridad establecido en la planeación. Se comenzó "
                      "con los módulos críticos (autenticación y base de datos), progresando "
                      "hacia los módulos de mayor complejidad (Stripe, webhooks y notificaciones push)."))
        setp(p, 401, "2.2.1 Implementación de módulos por prioridad")
        setp(p, 402, ("Se implementaron los módulos en el siguiente orden: (1) Configuración de "
                      "Supabase y políticas RLS; (2) Autenticación de usuarios con JWT; "
                      "(3) Catálogo de servicios y dashboard; (4) Integración de Stripe Checkout; "
                      "(5) Webhooks de confirmación transaccional; (6) Notificaciones push; "
                      "(7) Módulo de soporte con IA; (8) Panel administrativo."))
        setp(p, 404, "2.2.2 Integración de servicios externos")
        setp(p, 405, ("La integración de Supabase se realizó mediante su SDK oficial para Next.js. "
                      "La integración de Stripe se efectuó mediante la biblioteca stripe para "
                      "Node.js en API routes. La integración de Web Push utilizó la biblioteca "
                      "web-push con llaves VAPID generadas para producción."))

        setp(p, 407, "2.3 VERIFICAR y ACTUAR – Pruebas, Ajustes y Despliegue")
        setp(p, 408, ("En la fase de Verificar se ejecutaron pruebas funcionales para confirmar "
                      "el cumplimiento de cada criterio de aceptación: autenticación y autorización, "
                      "creación de sesión de Stripe Checkout, recepción y validación de webhooks, "
                      "envío y recepción de notificaciones push, y visualización de datos en "
                      "el dashboard."))
        setp(p, 409, ("Se verificó la compilación sin errores de TypeScript en modo estricto, "
                      "confirmando la coherencia de tipos entre todos los módulos del sistema. "
                      "Se probaron las políticas RLS mediante accesos cruzados entre usuarios, "
                      "confirmando el aislamiento de datos."))
        setp(p, 410, ("La fase de Actuar contempló los ajustes derivados de la verificación: "
                      "corrección de tipos TypeScript, optimización de consultas a Supabase, "
                      "ajuste de políticas RLS y configuración final de variables de entorno "
                      "para el despliegue en Vercel."))

    # =========================================================================
    # CAPÍTULO 3 – DESARROLLO
    # =========================================================================
    if todo or seccion == "cap3":
        print("✦ Escribiendo Capítulo 3 – Desarrollo (base)...")
        setp(p, 422, ("En este capítulo se documenta el desarrollo del sistema PagoIA, describiendo "
                      "en pasado e impersonal cada actividad realizada durante la residencia "
                      "profesional. Se incluyen evidencias técnicas mediante ilustraciones etiquetadas "
                      "y tablas de especificaciones que respaldan el trabajo ejecutado."))
        setp(p, 424, ("Las actividades se desarrollaron siguiendo el ciclo PHVA descrito en el "
                      "Capítulo 2, partiendo de la configuración de infraestructura hasta la "
                      "integración completa de todos los módulos funcionales del sistema."))
        setp(p, 425, ("Cada actividad generó evidencia técnica documentada mediante capturas de "
                      "pantalla (Ilustraciones), fragmentos de código representativos y resultados "
                      "de pruebas funcionales."))
        setp(p, 434, "3.1 Configuración de Infraestructura – Supabase y Base de Datos")
        setp(p, 435, "3.1.1 Modelado de entidades y despliegue de tablas")
        setp(p, 436, ("Se diseñaron y desplegaron en Supabase las tablas del sistema: profiles "
                      "(datos extendidos del usuario vinculados a auth.users), services (catálogo "
                      "de servicios disponibles), user_services (asociación usuario-servicio con "
                      "estado, monto y fecha de vencimiento), transactions (registro de pagos), "
                      "alerts (alertas por evento transaccional) y push_subscriptions "
                      "(suscripciones Web Push por usuario autenticado)."))
        setp(p, 437, "3.1.2 Implementación de políticas Row Level Security")
        setp(p, 438, ("Se configuraron políticas RLS en cada tabla para garantizar que los usuarios "
                      "solo accedieran a sus propios registros. Las políticas utilizan la función "
                      "auth.uid() para comparar el identificador del token JWT activo con el campo "
                      "user_id de cada fila. Se verificó el aislamiento de datos mediante pruebas "
                      "de acceso cruzado entre usuarios de prueba."))
        setp(p, 439, "3.2 Integración de Stripe – Checkout y Webhooks")
        setp(p, 440, ("Se integró Stripe Checkout en el flujo de pago: el usuario selecciona un "
                      "servicio desde el dashboard, el endpoint /api/payments/checkout-sessions "
                      "crea una sesión con los parámetros del servicio (monto en centavos MXN, "
                      "nombre y URLs de retorno), y el usuario completa el pago en el entorno "
                      "seguro de Stripe."))
        setp(p, 441, "3.2.1 Configuración de Stripe Checkout")
        setp(p, 442, ("Se configuraron los parámetros de sesión de checkout incluyendo monto, "
                      "moneda (MXN), nombre del servicio, URL de éxito y URL de cancelación. "
                      "La actualización definitiva del estado en la base de datos se realizó "
                      "mediante el webhook de confirmación, siguiendo las recomendaciones de "
                      "seguridad de Stripe."))
        setp(p, 443, ("Se gestionó la devolución del usuario desde Stripe hacia la aplicación, "
                      "validando el estado de la sesión de checkout y mostrando la confirmación "
                      "de pago correspondiente en el dashboard."))
        setp(p, 444, "3.2.2 Webhooks de confirmación transaccional")
        setp(p, 445, "3. Resumen de Actividades")
        setp(p, 446, ("Se configuró el endpoint /api/stripe/webhook para recibir eventos de Stripe. "
                      "Cada evento fue validado mediante la firma HMAC-SHA256 del encabezado "
                      "stripe-signature usando la clave secreta del webhook. Ante el evento "
                      "payment_intent.succeeded se actualizó el estado de la transacción en "
                      "Supabase y se generó una alerta para el usuario."))
        setp(p, 447, "3.3 Dashboard Operativo y Reportes")
        setp(p, 448, ("Se implementó el dashboard principal con indicadores de métricas "
                      "transaccionales: total de servicios activos, pagos del mes, balance estimado "
                      "y estado de vencimientos. Se incorporaron gráficos de ingresos por período "
                      "mediante Recharts."))
        setp(p, 449, ("Se desarrolló la sección de historial de transacciones con filtrado por "
                      "servicio, estado y período. El historial se actualiza en tiempo real mediante "
                      "suscripciones a canales de Supabase Realtime (postgres_changes)."))
        setp(p, 450, "3.4 Notificaciones Push y Sistema de Alertas")
        setp(p, 451, ("Se configuró el módulo de notificaciones push utilizando Web Push API con "
                      "llaves VAPID. El service worker (public/sw.js) se registró en el navegador "
                      "del usuario para habilitar la recepción de notificaciones en segundo plano."))
        setp(p, 452, ("Se verificó el envío de notificaciones push desde el servidor ante eventos "
                      "transaccionales relevantes (pago exitoso, alerta de vencimiento), confirmando "
                      "la recepción en el dispositivo del usuario."))
        setp(p, 453, ("En síntesis, las actividades de desarrollo cubrieron todas las fases del "
                      "ciclo PHVA: planeación de arquitectura y requerimientos; implementación de "
                      "módulos funcionales; verificación mediante pruebas de aceptación; y ajustes "
                      "previos al despliegue final en Vercel."))

    # =========================================================================
    # CAPÍTULO 4 – CONCLUSIONES
    # =========================================================================
    if todo or seccion == "cap4":
        print("✦ Escribiendo Capítulo 4 – Conclusiones...")
        setp(p, 465, ("El objetivo general del proyecto se logró mediante la implementación "
                      "incremental de los módulos funcionales del sistema PagoIA. Se desarrolló "
                      "una plataforma web integral capaz de centralizar autenticación, "
                      "administración de servicios, ejecución de pagos electrónicos seguros y "
                      "visualización de historial transaccional."))
        setp(p, 466, ("Los siete objetivos específicos se cumplieron: la interfaz web responsiva "
                      "fue implementada con Next.js 16.2.0, React 19 y Tailwind CSS; el backend "
                      "seguro opera mediante API routes con validación de autenticación; Supabase "
                      "utiliza políticas RLS; Stripe procesa transacciones con confirmación por "
                      "webhook; el módulo de IA provee soporte contextual; las pruebas verificaron "
                      "los flujos críticos; y la documentación técnica cubre arquitectura, modelo "
                      "de datos y guía de despliegue."))
        setp(p, 468, ("La evidencia funcional confirmó la correspondencia entre los requerimientos "
                      "definidos en análisis y los módulos desarrollados. La compilación sin errores "
                      "de TypeScript en modo estricto valida la calidad y coherencia del código."))
        setp(p, 469, ("La aplicación del ciclo PHVA permitió estructurar el desarrollo de forma "
                      "ordenada y verificable, con trazabilidad entre requerimientos, diseño, "
                      "implementación y validación."))
        setp(p, 473, ("Como trabajo futuro se recomienda incorporar analítica predictiva mediante "
                      "modelos de machine learning entrenados con el historial transaccional, e "
                      "integrar métodos de pago alternativos como OXXO Pay y transferencias SPEI."))
        setp(p, 474, ("Se recomienda fortalecer la estrategia de pruebas automatizadas end-to-end "
                      "con cobertura superior al 80%, ampliar el catálogo de servicios con nuevos "
                      "proveedores e implementar observabilidad con métricas y trazas distribuidas."))
        setp(p, 476, ("La residencia profesional en el proyecto PagoIA permitió consolidar "
                      "experiencia práctica en el ciclo completo de desarrollo de software: desde "
                      "el análisis de requerimientos y diseño de arquitectura hasta la implementación, "
                      "pruebas y despliegue en entorno de producción."))
        setp(p, 481, ("Durante la residencia se desarrollaron competencias de alto valor: diseño "
                      "de arquitecturas web por capas, modelado de datos relacional con RLS, "
                      "desarrollo Full-Stack con TypeScript en modo estricto, integración de APIs "
                      "externas certificadas, implementación de notificaciones push y despliegue "
                      "en plataformas cloud."))
        setp(p, 482, ("Competencias desarrolladas: dominio de TypeScript 5.7.3; desarrollo con "
                      "Next.js 16.2.0 y React 19; autenticación con Supabase Auth y JWT; "
                      "modelado con políticas RLS en PostgreSQL; integración de Stripe Checkout "
                      "y webhooks HMAC-SHA256; configuración de Web Push API con VAPID; "
                      "despliegue en Vercel con CI/CD."))

    # =========================================================================
    # REFERENCIAS APA
    # =========================================================================
    if todo or seccion == "referencias":
        print("✦ Escribiendo Referencias APA...")
        setp(p, 490, "REFERENCIAS")
        refs = [
            "1.  Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner's approach (9th ed.). McGraw-Hill.",
            "2.  Sommerville, I. (2016). Software engineering (10th ed.). Pearson.",
            "3.  Martin, R. C. (2009). Clean code: A handbook of agile software craftsmanship. Prentice Hall.",
            "4.  Bass, L., Clements, P., & Kazman, R. (2022). Software architecture in practice (4th ed.). Addison-Wesley.",
            "5.  Fowler, M. (2018). Refactoring: Improving the design of existing code (2nd ed.). Addison-Wesley.",
            "6.  Russell, S., & Norvig, P. (2021). Artificial intelligence: A modern approach (4th ed.). Pearson.",
            "7.  Laudon, K. C., & Laudon, J. P. (2020). Management information systems (16th ed.). Pearson.",
            "8.  Tanenbaum, A. S., & Wetherall, D. (2021). Computer networks (6th ed.). Pearson.",
            "9.  Stallings, W. (2020). Cryptography and network security (8th ed.). Pearson.",
            "10. Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design patterns. Addison-Wesley.",
            "11. Next.js. (2024). Next.js documentation. Vercel. Recuperado de https://nextjs.org/docs",
            "12. React. (2024). React documentation. Meta Platforms. Recuperado de https://react.dev",
            "13. TypeScript. (2024). TypeScript documentation. Microsoft. Recuperado de https://www.typescriptlang.org/docs",
            "14. Supabase. (2024). Supabase documentation. Supabase Inc. Recuperado de https://supabase.com/docs",
            "15. Stripe. (2024). Stripe documentation. Stripe Inc. Recuperado de https://stripe.com/docs",
            "16. Mozilla Developer Network. (2024). Web Push API. MDN. Recuperado de https://developer.mozilla.org/en-US/docs/Web/API/Push_API",
            "17. OWASP Foundation. (2024). OWASP Top Ten. Recuperado de https://owasp.org/www-project-top-ten",
            "18. Banco Interamericano de Desarrollo. (2022). Fintech en América Latina y el Caribe. Recuperado de https://publications.iadb.org",
            "19. Banco de México. (2023). Reporte del sistema financiero 2023. Recuperado de https://www.banxico.org.mx",
            "20. Node.js. (2024). Node.js documentation. OpenJS Foundation. Recuperado de https://nodejs.org/en/docs",
        ]
        base = 491
        for i, ref in enumerate(refs):
            idx = base + i * 2
            if idx < len(p):
                setp(p, idx, ref)

    # ─── Guardar versión base ─────────────────────────────────────────────────
    doc.save(str(SALIDA))
    print("  → Versión base guardada.")

    # =========================================================================
    # SEGUNDA PASADA: insertar Tablas e Ilustraciones al final (Cap 3)
    # =========================================================================
    if todo or seccion == "cap3":
        print("✦ Insertando Tablas e Ilustraciones (Capítulo 3 – segunda pasada)...")
        doc2 = Document(str(SALIDA))

        # ── Tabla 1: Stack tecnológico ─────────────────────────────────────
        add_tech_table(
            doc2,
            table_number=1,
            title="Stack tecnológico del sistema PagoIA",
            headers=["Tecnología", "Versión", "Rol en el sistema", "Licencia"],
            rows_data=[
                ["Next.js",      "16.2.0",  "Framework web – renderizado híbrido y API routes",           "MIT"],
                ["React",        "19.0.0",  "Biblioteca de componentes de interfaz reactiva",              "MIT"],
                ["TypeScript",   "5.7.3",   "Tipado estático y calidad del código en modo estricto",       "Apache 2.0"],
                ["Tailwind CSS", "3.4.x",   "Sistema de diseño utilitario y responsivo",                   "MIT"],
                ["Supabase",     "2.x",     "Autenticación JWT, PostgreSQL en la nube y RLS",              "Apache 2.0"],
                ["Stripe",       "16.x",    "Pasarela PCI-DSS – Checkout y Webhooks HMAC-SHA256",          "Propietaria"],
                ["Web Push API", "W3C",     "Notificaciones push con VAPID al dispositivo del usuario",    "W3C"],
                ["Vercel",       "—",       "Despliegue CI/CD, variables de entorno y edge network",       "Propietaria"],
                ["Node.js",      "22.x",    "Entorno de ejecución del servidor para API routes",           "MIT"],
            ]
        )

        # ── Tabla 2: Modelo de datos ───────────────────────────────────────
        add_tech_table(
            doc2,
            table_number=2,
            title="Entidades del modelo de datos en Supabase",
            headers=["Tabla", "Campos principales", "Propósito", "RLS"],
            rows_data=[
                ["profiles",          "id, full_name, avatar_url",                      "Datos extendidos del usuario",                    "Sí"],
                ["services",          "id, name, type, provider",                        "Catálogo de servicios disponibles",                "No"],
                ["user_services",     "id, user_id, service_id, status, due_date",       "Asociación usuario-servicio",                      "Sí"],
                ["transactions",      "id, user_id, amount, status, receipt_number",     "Registro de pagos y estado transaccional",         "Sí"],
                ["alerts",            "id, user_id, title, message, is_read",            "Alertas por eventos transaccionales",              "Sí"],
                ["push_subscriptions","id, user_id, endpoint, p256dh, auth",             "Suscripciones Web Push API",                       "Sí"],
            ]
        )

        # ── Tabla 3: Endpoints de API ──────────────────────────────────────
        add_tech_table(
            doc2,
            table_number=3,
            title="Endpoints de API implementados en el sistema PagoIA",
            headers=["Endpoint", "Método", "Descripción", "Auth"],
            rows_data=[
                ["/api/payments/checkout-sessions", "POST", "Crea sesión de Stripe Checkout",                  "JWT"],
                ["/api/stripe/webhook",             "POST", "Recibe y valida eventos Stripe (HMAC-SHA256)",     "Firma"],
                ["/api/notifications/subscribe",    "POST", "Registra suscripción Web Push del usuario",       "JWT"],
                ["/api/notifications/send",         "POST", "Envía notificación push al usuario",              "JWT"],
                ["/api/reports/monthly",            "GET",  "Métricas de transacciones del mes activo",        "JWT"],
                ["/api/ai-support",                 "POST", "Consultas al módulo de soporte con IA",           "JWT"],
                ["/api/vouchers/[id]",              "GET",  "Genera comprobante PDF de una transacción",       "JWT"],
            ]
        )

        # ── Ilustración 12: Esquema de tablas ─────────────────────────────
        add_illustration_box(
            doc2, 12,
            "Esquema de tablas en Supabase Table Editor",
            ("La ilustración muestra el modelo de datos desplegado en Supabase, incluyendo "
             "las seis tablas principales del sistema y sus relaciones de clave foránea hacia "
             "auth.users. Se observa la estructura de cada entidad con sus campos tipados, "
             "reflejando las decisiones de diseño tomadas en la fase Planear del ciclo PHVA. "
             "Las tablas con RLS habilitado garantizan que solo el usuario propietario de los "
             "registros pueda acceder o modificar sus datos.")
        )

        # ── Ilustración 11: Políticas RLS ─────────────────────────────────
        add_illustration_box(
            doc2, 11,
            "Políticas Row Level Security configuradas en Supabase",
            ("La ilustración presenta el panel de políticas RLS de Supabase para la tabla "
             "transactions. Se visualizan las políticas de lectura (SELECT) y escritura "
             "(INSERT, UPDATE) que utilizan auth.uid() para verificar que el user_id del "
             "registro coincida con el usuario autenticado. Esta configuración garantiza "
             "aislamiento total de datos entre usuarios sin necesidad de lógica adicional "
             "en la capa de aplicación Next.js.")
        )

        # ── Ilustración 1: Login ───────────────────────────────────────────
        add_illustration_box(
            doc2, 1,
            "Módulo de autenticación – pantalla de inicio de sesión de PagoIA",
            ("La ilustración muestra la pantalla de inicio de sesión del sistema PagoIA. "
             "El formulario permite autenticación mediante correo electrónico y contraseña, "
             "con validación del lado del cliente y mensajes de error contextuales. "
             "La autenticación es gestionada por Supabase Auth, que emite un token JWT "
             "al completar el proceso y redirige al usuario al dashboard.")
        )

        # ── Ilustración 6: Stripe Checkout ────────────────────────────────
        add_illustration_box(
            doc2, 6,
            "Flujo de Stripe Checkout – selección de servicio e inicio del proceso de pago",
            ("La ilustración presenta el flujo de pago iniciado desde el dashboard de PagoIA. "
             "Al seleccionar un servicio y presionar el botón de pago, el sistema llama al "
             "endpoint /api/payments/checkout-sessions, que crea la sesión en Stripe con el "
             "monto en centavos (MXN), nombre del servicio y URLs de retorno. El usuario es "
             "redirigido a la interfaz segura de Stripe, garantizando que los datos de tarjeta "
             "nunca pasen por los servidores de la aplicación PagoIA.")
        )

        # ── Ilustración 7: Confirmación de pago ───────────────────────────
        add_illustration_box(
            doc2, 7,
            "Confirmación de pago exitoso y actualización de estado en el dashboard",
            ("Tras completar el pago en Stripe, el usuario es redirigido a la URL de éxito "
             "(/dashboard?payment=success). El dashboard muestra un mensaje de confirmación "
             "y actualiza el estado del servicio. La actualización definitiva en Supabase "
             "ocurre mediante el webhook payment_intent.succeeded, que registra la transacción "
             "con estado 'completed', monto, fecha y número de comprobante.")
        )

        # ── Ilustración 8: Webhooks Stripe ────────────────────────────────
        add_illustration_box(
            doc2, 8,
            "Panel de configuración de Webhooks en Stripe Dashboard",
            ("La ilustración muestra la configuración del endpoint de webhook en el panel "
             "de Stripe, incluyendo la URL del endpoint (/api/stripe/webhook), los eventos "
             "suscritos (payment_intent.succeeded, payment_intent.payment_failed) y el "
             "historial de eventos recibidos con código de respuesta HTTP 200, confirmando "
             "la correcta recepción y validación de eventos mediante firma HMAC-SHA256.")
        )

        # ── Ilustración 3: Dashboard ───────────────────────────────────────
        add_illustration_box(
            doc2, 3,
            "Dashboard principal de PagoIA con indicadores de servicios y métricas transaccionales",
            ("El dashboard principal integra los módulos operativos del sistema: indicadores "
             "de métricas (servicios activos, pagos del mes, balance estimado), gráfico de "
             "ingresos por período mediante Recharts, lista de servicios con estado de "
             "vencimiento y acceso directo al módulo de pago. Los datos se obtienen de "
             "Supabase Realtime, actualizándose automáticamente ante nuevas transacciones.")
        )

        # ── Ilustración 4: Historial ───────────────────────────────────────
        add_illustration_box(
            doc2, 4,
            "Historial de transacciones con filtrado por servicio, estado y período",
            ("La sección de historial de transacciones presenta el registro completo de "
             "pagos con filtrado por tipo de servicio, estado (completado, pendiente, fallido) "
             "y rango de fechas. Cada registro muestra nombre del servicio, proveedor, monto, "
             "fecha y número de comprobante. Los datos se actualizan en tiempo real "
             "mediante suscripciones Supabase Realtime.")
        )

        # ── Ilustración 9: Notificaciones push ────────────────────────────
        add_illustration_box(
            doc2, 9,
            "Suscripción a notificaciones push desde el navegador",
            ("La ilustración muestra el diálogo del navegador solicitando permiso para "
             "enviar notificaciones push al usuario. Al aceptar, el service worker registra "
             "la suscripción con las llaves VAPID y la envía al endpoint "
             "/api/notifications/subscribe, donde se almacena en push_subscriptions de "
             "Supabase. Este flujo habilita notificaciones incluso cuando la aplicación "
             "no está abierta en el navegador.")
        )

        # ── Ilustración 10: Centro de notificaciones ──────────────────────
        add_illustration_box(
            doc2, 10,
            "Centro de notificaciones y alertas del dashboard de PagoIA",
            ("El centro de notificaciones presenta las alertas del usuario con indicadores "
             "visuales de tipo (éxito, advertencia, error) y estado de lectura. Cada alerta "
             "incluye título, mensaje descriptivo y marca de tiempo. Las alertas se generan "
             "automáticamente ante eventos transaccionales: pago exitoso, fallo en cobro, "
             "disputa recibida o vencimiento próximo.")
        )

        # ── Ilustración 14: Módulo de IA ───────────────────────────────────
        add_illustration_box(
            doc2, 14,
            "Módulo de soporte con Inteligencia Artificial integrado en el dashboard",
            ("La ilustración presenta el panel flotante de soporte con IA, accesible desde "
             "cualquier sección del dashboard. El módulo permite consultas en lenguaje natural "
             "sobre historial de pagos, estado de servicios y recomendaciones financieras. "
             "Las respuestas se generan con acceso controlado al contexto transaccional del "
             "usuario autenticado, validando el JWT en el endpoint /api/ai-support.")
        )

        # ── Ilustración 15: Vercel deploy ─────────────────────────────────
        add_illustration_box(
            doc2, 15,
            "Despliegue exitoso en Vercel con configuración de variables de entorno de producción",
            ("La ilustración muestra el panel de Vercel con el despliegue exitoso del "
             "proyecto PagoIA: estado del build (éxito), dominio de producción, rama de "
             "deployment (main) y la sección de variables de entorno donde se configuran "
             "las credenciales de producción (Supabase URL, Supabase Anon Key, "
             "Stripe Secret Key, Stripe Webhook Secret y llaves VAPID). El pipeline CI/CD "
             "ejecuta el build automáticamente ante cada push a la rama principal.")
        )

        doc2.save(str(SALIDA))
        print("  → Tablas e Ilustraciones insertadas.")

    # ─── Estadísticas finales ─────────────────────────────────────────────────
    doc_final = Document(str(SALIDA))
    words = sum(len(par.text.split()) for par in doc_final.paragraphs if par.text.strip())
    tables_count = len(doc_final.tables)
    print()
    print(f"✅  Documento generado : {SALIDA.name}")
    print(f"    Párrafos totales   : {len(doc_final.paragraphs)}")
    print(f"    Tablas insertadas  : {tables_count}")
    print(f"    Palabras aprox.    : {words}")
    print(f"    Páginas estimadas  : ~{words // 350} páginas")
    print()
    print("RECUADROS PARA CAPTURAS (Ilustraciones a reemplazar con capturas reales):")
    ilustraciones = [
        " 1 – Módulo de autenticación / Login",
        " 3 – Dashboard principal con métricas",
        " 4 – Historial de transacciones",
        " 6 – Flujo de Stripe Checkout",
        " 7 – Confirmación de pago exitoso",
        " 8 – Panel de Webhooks en Stripe Dashboard",
        " 9 – Suscripción a notificaciones push",
        "10 – Centro de notificaciones y alertas",
        "11 – Políticas RLS en Supabase",
        "12 – Esquema de tablas en Supabase",
        "14 – Módulo de soporte con IA",
        "15 – Despliegue exitoso en Vercel",
    ]
    for il in ilustraciones:
        print(f"    Ilustración {il}")


if __name__ == "__main__":
    seccion = None
    if "--seccion" in sys.argv:
        idx = sys.argv.index("--seccion")
        if idx + 1 < len(sys.argv):
            seccion = sys.argv[idx + 1].lower()
            print(f"Modo: actualizar sección '{seccion}'")
        else:
            print("ERROR: falta nombre de sección después de --seccion")
            sys.exit(1)
    else:
        print("Modo: generar documento COMPLETO")

    run(seccion)
